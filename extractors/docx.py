from pathlib import Path

from docx import Document


def extract_docx(path: Path, max_chars: int) -> str:
    doc = Document(str(path))

    parts = [f"[DOCX] {path.name}"]
    char_count = 0
    truncated = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        line = f"[{style_name}] {text}" if style_name.lower().startswith("heading") else text

        if char_count + len(line) > max_chars:
            truncated = True
            break
        parts.append(line)
        char_count += len(line)

    if not truncated:
        for ti, table in enumerate(doc.tables, start=1):
            table_lines = [f"\n--- Table {ti} ---"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            table_text = "\n".join(table_lines)

            if char_count + len(table_text) > max_chars:
                truncated = True
                break
            parts.append(table_text)
            char_count += len(table_text)

    if truncated:
        parts.append(f"\n[문자 수 제한({max_chars})으로 이후 내용 생략]")

    return "\n".join(parts)
